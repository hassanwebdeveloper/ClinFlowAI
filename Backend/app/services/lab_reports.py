"""Classify lab uploads and extract structured details (vision vs text-only path)."""

from __future__ import annotations

import json
import logging
import mimetypes
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import fitz

from app.services.together import (
    extract_lab_report_with_vl,
    extract_lab_reports_from_transcript as _llm_lab_reports_from_transcript,
    infer_abnormal_flags_for_analytes,
    normalize_lab_report_from_text,
)

logger = logging.getLogger(__name__)

MIN_TEXT_CHARS_TO_SKIP_VL = 80
MAX_PDF_PAGES_FOR_VL = 8


_VALID_ABNORMAL_FLAGS = {"", "H", "L", "critical"}


def _full_lab_text_for_judge(suggested_test_name: str, details_body: str) -> str:
    """Narrative block passed to the abnormal-flag judge (panel name + full extracted summary text)."""
    name = (suggested_test_name or "").strip()
    body = (details_body or "").strip()
    if name and body:
        return f"Panel / test: {name}\n\n{body}"
    return name or body


def _coerce_number(val: Any) -> float | None:
    """Best-effort numeric coercion for LLM-emitted analyte fields.

    Accepts numbers, numeric strings, and strings containing one numeric token.
    Returns None for anything else (so the chart simply skips that point).
    """
    if val is None:
        return None
    if isinstance(val, bool):
        return None
    if isinstance(val, (int, float)):
        if val != val:  # NaN
            return None
        return float(val)
    if isinstance(val, str):
        s = val.strip().replace(",", "")
        if not s:
            return None
        try:
            return float(s)
        except ValueError:
            return None
    return None


def _normalize_analyte(raw: Any) -> dict[str, Any] | None:
    if not isinstance(raw, dict):
        return None
    name = str(raw.get("name", "")).strip()
    if not name:
        return None
    value = _coerce_number(raw.get("value"))
    unit = str(raw.get("unit", "") or "").strip()
    ref_low = _coerce_number(raw.get("ref_low"))
    ref_high = _coerce_number(raw.get("ref_high"))
    flag_raw = str(raw.get("abnormal_flag", "") or "").strip()
    flag = flag_raw if flag_raw in _VALID_ABNORMAL_FLAGS else ""
    qual = str(raw.get("qualitative_value") or "").strip()
    if value is None and not flag and not qual:
        # Skip rows with no numeric value, no flag, and no qualitative text.
        return None
    row: dict[str, Any] = {
        "name": name,
        "value": value,
        "unit": unit,
        "ref_low": ref_low,
        "ref_high": ref_high,
        "abnormal_flag": flag,
    }
    if qual:
        row["qualitative_value"] = qual
    return row


def _parse_analytes_payload(raw_json: str) -> list[dict[str, Any]]:
    raw_json = (raw_json or "").strip()
    if not raw_json:
        return []
    try:
        parsed = json.loads(raw_json)
    except json.JSONDecodeError:
        return []
    items: list[Any]
    if isinstance(parsed, dict):
        candidate = parsed.get("analytes")
        items = candidate if isinstance(candidate, list) else []
    elif isinstance(parsed, list):
        items = parsed
    else:
        return []
    out: list[dict[str, Any]] = []
    for it in items:
        norm = _normalize_analyte(it)
        if norm is not None:
            out.append(norm)
    return out


async def _apply_inferred_abnormal_flags(
    analytes: list[dict[str, Any]],
    full_lab_report_text: str = "",
) -> None:
    """Second LLM pass: set abnormal_flag using full report text plus structured rows."""
    if not analytes:
        return
    try:
        flags = await infer_abnormal_flags_for_analytes(analytes, full_lab_report_text)
        for i, a in enumerate(analytes):
            if i < len(flags) and flags[i] in _VALID_ABNORMAL_FLAGS:
                a["abnormal_flag"] = flags[i]
    except Exception:
        logger.exception("infer abnormal flags failed; leaving flags from extraction")


def _merge_analytes(*groups: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Union analytes across groups, de-duped case-insensitively by name (first wins)."""
    seen: set[str] = set()
    out: list[dict[str, Any]] = []
    for g in groups:
        for a in g or []:
            key = (a.get("name") or "").strip().lower()
            if not key or key in seen:
                continue
            seen.add(key)
            out.append(a)
    return out


def split_lab_headers_from_extraction(
    details: str,
) -> tuple[str, str, list[dict[str, Any]], str]:
    """
    Parse LAB_TEST_NAME, LAB_TEST_PATTERN, and LAB_ANALYTES_JSON; strip header lines from
    the body shown in UI. Returns (lab_test_name, lab_test_pattern, analytes, body).
    Pattern is stored in DB only (e.g. [monitoring]).
    """
    if not details or not details.strip():
        return "", "", [], ""
    lines = details.split("\n")
    name = ""
    pattern = ""
    analytes_raw = ""
    name_seen = False
    pattern_seen = False
    analytes_seen = False
    out: list[str] = []
    for line in lines:
        stripped = line.strip()
        upper = stripped.upper()
        if upper.startswith("LAB_TEST_NAME:") or upper.startswith("PANEL_NAME:"):
            if not name_seen:
                raw = line.split(":", 1)[1].strip() if ":" in line else ""
                if raw and raw.upper() not in ("UNKNOWN", "N/A", "NONE", "UNCLEAR"):
                    name = raw
                name_seen = True
            continue
        if upper.startswith("LAB_TEST_PATTERN:"):
            if not pattern_seen:
                pattern = line.split(":", 1)[1].strip() if ":" in line else ""
                pattern_seen = True
            continue
        if upper.startswith("LAB_ANALYTES_JSON:"):
            if not analytes_seen:
                analytes_raw = line.split(":", 1)[1].strip() if ":" in line else ""
                analytes_seen = True
            continue
        out.append(line)
    analytes = _parse_analytes_payload(analytes_raw)
    return name, pattern, analytes, "\n".join(out).strip()


def _pdf_extract_text(path: str) -> str:
    doc = fitz.open(path)
    try:
        parts: list[str] = []
        for i in range(doc.page_count):
            parts.append(doc[i].get_text())
        return "\n".join(parts).strip()
    finally:
        doc.close()


def _pdf_pages_as_png(path: str, max_pages: int) -> list[bytes]:
    doc = fitz.open(path)
    try:
        out: list[bytes] = []
        n = min(doc.page_count, max_pages)
        mat = fitz.Matrix(2, 2)
        for i in range(n):
            pix = doc[i].get_pixmap(matrix=mat, alpha=False)
            out.append(pix.tobytes("png"))
        return out
    finally:
        doc.close()


def _docx_extract_text(path: str) -> str:
    from docx import Document

    doc = Document(path)
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _resolve_mime(filename: str, content_type: str | None) -> str:
    ct = (content_type or "").split(";")[0].strip().lower()
    if ct and ct != "application/octet-stream":
        return ct
    guessed, _ = mimetypes.guess_type(filename)
    return (guessed or "application/octet-stream").lower()


_IMAGE_EXT = {
    ".jpg",
    ".jpeg",
    ".png",
    ".webp",
    ".gif",
    ".bmp",
    ".tif",
    ".tiff",
    ".heic",
    ".heif",
}


def _looks_like_image_file(path: Path, original_filename: str, content_type: str | None) -> bool:
    mime = _resolve_mime(original_filename, content_type)
    ext = path.suffix.lower()
    return mime.startswith("image/") or ext in _IMAGE_EXT


async def extract_lab_from_image_group(
    items: list[tuple[str, str, str | None]],
) -> tuple[str, str, str, str, list[dict[str, Any]]]:
    """
    One logical report from multiple photos (e.g. multi-page lab result photographed in parts).
    Each image is passed through VL; outputs (including analytes JSON) are merged.
    """
    if not items:
        raise ValueError("No lab images")
    if len(items) == 1:
        return await extract_lab_from_saved_file(items[0][0], items[0][1], items[0][2])
    per_chunk_analytes: list[list[dict[str, Any]]] = []
    name = ""
    pattern = ""
    body_parts: list[str] = []
    for i, (path, fname, ctype) in enumerate(items):
        p = Path(path)
        if not _looks_like_image_file(p, fname, ctype):
            raise ValueError(
                f"Merged lab report must be images only; put {fname!r} in a separate upload"
            )
        mime = _resolve_mime(fname, ctype)
        img_mime = mime if mime.startswith("image/") else "image/jpeg"
        data = p.read_bytes()
        raw = (await extract_lab_report_with_vl(data, img_mime)).strip()
        label = fname or f"part-{i + 1}"
        n_i, p_i, a_i, b_i = split_lab_headers_from_extraction(raw)
        if not name:
            name = n_i
        if not pattern:
            pattern = p_i
        per_chunk_analytes.append(a_i)
        body_parts.append(f"--- Photo {i + 1} of {len(items)} ({label}) ---\n{b_i}")
    analytes = _merge_analytes(*per_chunk_analytes)
    body_merged = "\n\n".join(body_parts).strip()
    await _apply_inferred_abnormal_flags(analytes, _full_lab_text_for_judge(name, body_merged))
    return body_merged, "vl", name, pattern, analytes


async def extract_lab_from_saved_file(
    saved_path: str,
    original_filename: str,
    content_type: str | None,
) -> tuple[str, str, str, str, list[dict[str, Any]]]:
    """
    Returns (details_text, extraction_method, suggested_lab_test_name, lab_test_pattern, analytes).
    lab_test_pattern is e.g. [one-time] or [monitoring]; stored in DB, omitted from UI.
    analytes is the structured list parsed from the LAB_ANALYTES_JSON header (may be empty).
    """
    path = Path(saved_path)
    ext = path.suffix.lower()
    mime = _resolve_mime(original_filename, content_type)

    if mime.startswith("image/") or ext in {
        ".jpg",
        ".jpeg",
        ".png",
        ".webp",
        ".gif",
        ".bmp",
        ".tif",
        ".tiff",
        ".heic",
        ".heif",
    }:
        data = path.read_bytes()
        img_mime = mime if mime.startswith("image/") else "image/jpeg"
        raw = (await extract_lab_report_with_vl(data, img_mime)).strip()
        sugg, patt, analytes, body = split_lab_headers_from_extraction(raw)
        await _apply_inferred_abnormal_flags(analytes, _full_lab_text_for_judge(sugg, body))
        return body, "vl", sugg, patt, analytes

    if mime in {"text/plain", "text/csv"} or ext in {".txt", ".csv"}:
        raw_file = path.read_text(encoding="utf-8", errors="replace").strip()
        if len(raw_file) < 3:
            raise ValueError("Lab text file is empty")
        raw = (await normalize_lab_report_from_text(raw_file)).strip()
        sugg, patt, analytes, body = split_lab_headers_from_extraction(raw)
        await _apply_inferred_abnormal_flags(analytes, _full_lab_text_for_judge(sugg, body))
        return body, "text", sugg, patt, analytes

    if ext == ".docx" or mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        raw_doc = _docx_extract_text(str(path))
        if len(raw_doc) < MIN_TEXT_CHARS_TO_SKIP_VL:
            raise ValueError("Could not extract enough text from the Word document")
        raw = (await normalize_lab_report_from_text(raw_doc)).strip()
        sugg, patt, analytes, body = split_lab_headers_from_extraction(raw)
        await _apply_inferred_abnormal_flags(analytes, _full_lab_text_for_judge(sugg, body))
        return body, "text", sugg, patt, analytes

    if ext == ".pdf" or mime == "application/pdf":
        text = _pdf_extract_text(str(path))
        if len(text) >= MIN_TEXT_CHARS_TO_SKIP_VL:
            raw = (await normalize_lab_report_from_text(text[:50_000])).strip()
            sugg, patt, analytes, body = split_lab_headers_from_extraction(raw)
            await _apply_inferred_abnormal_flags(analytes, _full_lab_text_for_judge(sugg, body))
            return body, "text", sugg, patt, analytes
        pngs = _pdf_pages_as_png(str(path), MAX_PDF_PAGES_FOR_VL)
        if not pngs:
            raise ValueError("PDF has no pages to read")
        per_page_analytes: list[list[dict[str, Any]]] = []
        name = ""
        pattern = ""
        body_parts: list[str] = []
        for i, png in enumerate(pngs):
            part = (await extract_lab_report_with_vl(png, "image/png")).strip()
            n_i, p_i, a_i, b_i = split_lab_headers_from_extraction(part)
            if not name:
                name = n_i
            if not pattern:
                pattern = p_i
            per_page_analytes.append(a_i)
            body_parts.append(f"--- PDF page {i + 1} ---\n{b_i}")
        analytes = _merge_analytes(*per_page_analytes)
        body_merged = "\n\n".join(body_parts).strip()
        await _apply_inferred_abnormal_flags(analytes, _full_lab_text_for_judge(name, body_merged))
        return body_merged, "vl", name, pattern, analytes

    raise ValueError(f"Unsupported lab file type: {original_filename}")


# Spoken-test aliases → canonical display name for transcript-derived lab items.
_TRANSCRIPT_TEST_ALIASES: dict[str, str] = {
    "hba1c": "HbA1c",
    "hb a1c": "HbA1c",
    "hemoglobin a1c": "HbA1c",
    "a1c": "HbA1c",
    "fasting sugar": "Fasting Blood Sugar",
    "fasting blood sugar": "Fasting Blood Sugar",
    "fasting glucose": "Fasting Blood Sugar",
    "fbs": "Fasting Blood Sugar",
    "blood sugar fasting": "Fasting Blood Sugar",
    "creatinine": "Creatinine",
    "serum creatinine": "Creatinine",
    "cbc": "CBC",
    "complete blood count": "CBC",
    "lipid profile": "Lipid Profile",
    "lipid panel": "Lipid Profile",
    "tsh": "TSH",
    "thyroid stimulating hormone": "TSH",
}

# Ordered patterns: first match in a details line wins for inferring test_name.
_TRANSCRIPT_LINE_TEST_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"\b(?:hb\s*a1c|hemoglobin\s*a1c|a1c)\b", re.I), "HbA1c"),
    (re.compile(r"\b(?:fasting\s*(?:blood\s*)?sugar|fasting\s*glucose|fbs)\b", re.I), "Fasting Blood Sugar"),
    (re.compile(r"\b(?:serum\s*)?creatinine\b", re.I), "Creatinine"),
    (re.compile(r"\b(?:complete\s*blood\s*count|cbc)\b", re.I), "CBC"),
    (re.compile(r"\b(?:lipid\s*(?:profile|panel)|cholesterol\s*panel)\b", re.I), "Lipid Profile"),
    (re.compile(r"\b(?:thyroid\s*stimulating\s*hormone|tsh)\b", re.I), "TSH"),
]


def _canonical_transcript_test_name(raw: str) -> str:
    """Normalize spoken/LLM test names (strip list prefixes, apply aliases)."""
    s = re.sub(r"^\d+\.\s*", "", (raw or "").strip())
    if not s:
        return ""
    return _TRANSCRIPT_TEST_ALIASES.get(s.lower(), s)


def _infer_test_name_from_text(text: str) -> str:
    """Best-effort test name from a details line or analyte name."""
    t = (text or "").strip()
    if not t:
        return ""
    for pat, name in _TRANSCRIPT_LINE_TEST_PATTERNS:
        if pat.search(t):
            return name
    key = t.lower()
    if key in _TRANSCRIPT_TEST_ALIASES:
        return _TRANSCRIPT_TEST_ALIASES[key]
    # "Creatinine was 1.0" → take token before was/is/result/came
    m = re.match(
        r"^([A-Za-z][A-Za-z0-9\s\-/]{0,40}?)\s+(?:was|is|are|result|came\s+back)\b",
        t,
        re.I,
    )
    if m:
        return _canonical_transcript_test_name(m.group(1).strip())
    return _canonical_transcript_test_name(t.split()[0] if t.split() else "")


def _detail_bullets(details: str) -> list[str]:
    """Split a bundled details field into non-empty bullet lines."""
    bullets: list[str] = []
    for line in (details or "").split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        while stripped and stripped[0] in "-•*":
            stripped = stripped.lstrip("-•*").strip()
        if stripped:
            bullets.append(stripped)
    return bullets


def _analyte_belongs_to_test(analyte_name: str, test_name: str) -> bool:
    """Heuristic: does this analyte row belong under the given panel test_name?"""
    a = (analyte_name or "").strip().lower()
    t = _canonical_transcript_test_name(test_name).lower()
    if not a or not t:
        return False
    if t in a or a in t:
        return True
    if t == "hba1c" and ("a1c" in a or "hba1c" in a or "hemoglobin a1c" in a):
        return True
    if t == "fasting blood sugar" and any(
        x in a for x in ("glucose", "sugar", "fbs", "fasting")
    ):
        return True
    if t == "creatinine" and "creatinine" in a:
        return True
    if t == "cbc" and any(
        x in a
        for x in (
            "wbc",
            "rbc",
            "hemoglobin",
            "hgb",
            "platelet",
            "plt",
            "mcv",
            "mch",
            "hematocrit",
            "hct",
            "cbc",
        )
    ):
        return True
    return _infer_test_name_from_text(analyte_name).lower() == t


def _split_bundled_transcript_lab_item(item: dict[str, Any]) -> list[dict[str, Any]]:
    """Split one LLM item that incorrectly merged multiple spoken lab results."""
    pattern = (item.get("lab_test_pattern") or "").strip()
    details = (item.get("details") or "").strip()
    analytes_raw = item.get("analytes") or []
    analytes = [a for a in analytes_raw if isinstance(a, dict)]

    # Multiple structured analytes → one lab report per analyte.
    if len(analytes) > 1:
        out: list[dict[str, Any]] = []
        for a in analytes:
            an_name = str(a.get("name", "")).strip()
            test = _infer_test_name_from_text(an_name) or _canonical_transcript_test_name(
                str(item.get("test_name", ""))
            )
            out.append({
                "test_name": test or an_name,
                "lab_test_pattern": pattern,
                "details": an_name,
                "analytes": [a],
            })
        return out

    bullets = _detail_bullets(details)
    if len(bullets) > 1:
        out = []
        for b in bullets:
            test = _infer_test_name_from_text(b)
            if not test:
                test = _canonical_transcript_test_name(str(item.get("test_name", "")))
            matched: list[dict[str, Any]] = []
            if len(analytes) == 1 and _analyte_belongs_to_test(
                str(analytes[0].get("name", "")), test
            ):
                matched = analytes
            out.append({
                "test_name": test or b[:48],
                "lab_test_pattern": pattern,
                "details": b,
                "analytes": matched,
            })
        return out

    name = _canonical_transcript_test_name(str(item.get("test_name", "")))
    return [{
        "test_name": name or str(item.get("test_name", "")).strip(),
        "lab_test_pattern": pattern,
        "details": details,
        "analytes": analytes,
    }]


def _split_bundled_transcript_lab_items(
    items: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    for item in items:
        if not isinstance(item, dict):
            continue
        out.extend(_split_bundled_transcript_lab_item(item))
    return out


async def extract_transcript_lab_reports(
    transcript: str,
    excluded_test_names: list[str] | None = None,
) -> list[dict[str, Any]]:
    """Extract lab reports the doctor verbally described in the visit transcript.

    Uploaded files take priority — anything in `excluded_test_names`
    (case-insensitively) is filtered out by both the prompt and a final pass here.
    Returns a list of dicts ready to be stored as `lab_reports` records:
        {test_name, lab_test_pattern, details, analytes}
    Analytes are normalized through the same `_normalize_analyte` helper used by
    the file-based pipeline, so charts/alerts treat them identically. Items left
    with neither a normalized analyte nor a non-empty details summary are dropped.
    """
    raw_items = _split_bundled_transcript_lab_items(
        await _llm_lab_reports_from_transcript(transcript, excluded_test_names)
    )
    if not raw_items:
        return []
    excluded_lower = {
        (n or "").strip().lower()
        for n in (excluded_test_names or [])
        if n and n.strip()
    }
    out: list[dict[str, Any]] = []
    for item in raw_items:
        name = (item.get("test_name") or "").strip()
        if not name or name.lower() in excluded_lower:
            continue
        analytes_norm: list[dict[str, Any]] = []
        for a in item.get("analytes") or []:
            norm = _normalize_analyte(a)
            if norm is not None:
                analytes_norm.append(norm)
        details = (item.get("details") or "").strip()
        if analytes_norm:
            await _apply_inferred_abnormal_flags(
                analytes_norm,
                _full_lab_text_for_judge(name, details),
            )
        if not analytes_norm and not details:
            continue
        out.append({
            "test_name": name,
            "lab_test_pattern": (item.get("lab_test_pattern") or "").strip(),
            "details": details,
            "analytes": analytes_norm,
        })
    return out


def transcript_lab_report_records(
    items: list[dict[str, Any]],
    visit_id: str,
    *,
    id_base_ms: int | None = None,
) -> list[dict[str, Any]]:
    """Build patient-level lab_reports rows for transcript-derived extractions."""
    recorded_at = datetime.now(timezone.utc).isoformat()
    base_lr = (
        id_base_ms
        if id_base_ms is not None
        else int(datetime.now(timezone.utc).timestamp() * 1000)
    )
    out: list[dict[str, Any]] = []
    for j, item in enumerate(items):
        out.append({
            "id": f"lr-{base_lr}-t{j}",
            "recorded_at": recorded_at,
            "filename": "",
            "extraction_method": "transcript",
            "details": item.get("details") or "",
            "test_name": item.get("test_name") or "",
            "lab_test_pattern": item.get("lab_test_pattern") or "",
            "visit_id": visit_id,
            "file_id": None,
            "file_url": None,
            "extra_file_ids": [],
            "extra_file_urls": [],
            "analytes": list(item.get("analytes") or []),
        })
    return out


async def refresh_transcript_lab_reports_for_visit(
    lab_reports: list[dict[str, Any]],
    visit_id: str,
    transcript: str,
) -> list[dict[str, Any]]:
    """Replace transcript-derived lab rows for one visit; keep uploads and other visits."""
    vid = (visit_id or "").strip()
    kept: list[dict[str, Any]] = []
    uploaded_names: list[str] = []
    prior_transcript: list[dict[str, Any]] = []
    for lr in lab_reports:
        if not isinstance(lr, dict):
            continue
        if (lr.get("visit_id") or "").strip() != vid:
            kept.append(lr)
            continue
        if (lr.get("extraction_method") or "").strip() == "transcript":
            prior_transcript.append(lr)
            continue
        kept.append(lr)
        name = (lr.get("test_name") or "").strip()
        if name:
            uploaded_names.append(name)

    try:
        extracted = await extract_transcript_lab_reports(transcript, uploaded_names)
    except Exception:
        logger.exception("transcript lab re-extraction failed for visit %s", vid)
        extracted = []

    if extracted:
        return kept + transcript_lab_report_records(extracted, vid)

    # LLM unavailable or returned nothing: split any bundled rows already on file.
    fallback: list[dict[str, Any]] = []
    for lr in prior_transcript:
        item = {
            "test_name": lr.get("test_name") or "",
            "lab_test_pattern": lr.get("lab_test_pattern") or "",
            "details": lr.get("details") or "",
            "analytes": list(lr.get("analytes") or []),
        }
        split_items = _split_bundled_transcript_lab_items([item])
        if len(split_items) <= 1:
            fallback.append(lr)
            continue
        base_ms = None
        lid = str(lr.get("id") or "")
        if lid.startswith("lr-") and "-t" in lid:
            try:
                base_ms = int(lid.split("-", 2)[1])
            except ValueError:
                base_ms = None
        fresh = transcript_lab_report_records(
            split_items, vid, id_base_ms=base_ms
        )
        recorded_at = lr.get("recorded_at")
        if recorded_at:
            for row in fresh:
                row["recorded_at"] = recorded_at
        fallback.extend(fresh)
    return kept + fallback
